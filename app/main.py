# from fastapi import FastAPI, HTTPException
# from fastapi.responses import JSONResponse
# from pydantic import BaseModel, Field
# from dotenv import load_dotenv
# from fastapi.middleware.cors import CORSMiddleware
# import os
# from PIL import Image, UnidentifiedImageError
# import numpy as np
# from paddleocr import PaddleOCR
# import cohere
# import fitz  # PyMuPDF
# from docx import Document
# import io
# import json
# import re
# import base64
# import time

# # Load environment variables
# load_dotenv()

# # Get the Cohere API key from environment variables
# api_key = os.getenv("COHERE_API_KEY")
# if not api_key:
#     raise ValueError("API key not found. Please set it in the .env file.")

# # Initialize the Cohere client
# co = cohere.Client(api_key)

# # Initialize PaddleOCR
# ocr = PaddleOCR(use_angle_cls=True, lang='en')

# app = FastAPI()

# # Add CORS middleware
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # Allow all origins, change this to specific origins as needed
#     allow_credentials=True,
#     allow_methods=["*"],  # Allow all methods
#     allow_headers=["*"],  # Allow all headers
# )

# class DocumentAttachment(BaseModel):
#     base64Data: str = Field(..., description="The base64-encoded content of the file")
#     type: str = Field(..., description="The MIME type of the file")

# class DocumentAttachments(BaseModel):
#     DocumentAttachments: list[DocumentAttachment]

# def extract_text_and_boxes_from_image(image_bytes: bytes):
#     try:
#         image = Image.open(io.BytesIO(image_bytes))
#         if image.mode in ["RGBA", "P"]:
#             image = image.convert("RGB")
#         image_array = np.array(image)
#         result = ocr.ocr(image_array, cls=True)
#         texts = [line[1][0] for line in result[0]]
#         return "\n".join(texts)  # Join texts with newline for better readability
#     except UnidentifiedImageError:
#         raise HTTPException(status_code=400, detail="Failed to identify the image file")

# def process_text_with_cohere(text, prompt):
#     response = co.generate(
#         model='command-xlarge-nightly',
#         prompt=f'{prompt} {text}',
#         max_tokens=500,
#         temperature=0.5,
#         stop_sequences=["--END--"]
#     )
#     # Log the raw response for debugging
#     print("Raw response from Cohere:", response)

#     # Check if the response contains generations and is not empty
#     if not response or not response.generations:
#         raise ValueError("Empty response or no generations returned from Cohere")

#     return response.generations[0].text.strip()

# def extract_text_from_pdf(file_bytes: bytes):
#     pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
#     texts = []
#     images = []

#     for page_num in range(len(pdf_document)):
#         page = pdf_document.load_page(page_num)
#         texts.append(page.get_text())

#         # Extract images from the page
#         image_list = page.get_images(full=True)
#         for img_index, img in enumerate(image_list):
#             xref = img[0]
#             base_image = pdf_document.extract_image(xref)
#             image_bytes = base_image["image"]
#             images.append(image_bytes)

#     return texts, images

# def extract_text_from_docx(file_bytes: bytes):
#     doc = Document(io.BytesIO(file_bytes))
#     text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
#     return [text]

# def clean_json_string(json_string: str) -> str:
#     # Remove any unwanted characters
#     json_string = re.sub(r'\n', '', json_string)
#     json_string = re.sub(r'\s+', ' ', json_string)
#     return json_string

# @app.post("/extract", response_model=dict)
# async def extract_data(document_attachments: DocumentAttachments):
#     results = []

#     try:
#         start_time = time.time()  # Record start time

#         for attachment in document_attachments.DocumentAttachments:
#             file_result = {}

#             # Decode the base64-encoded file content
#             try:
#                 file_bytes = base64.b64decode(attachment.base64Data)
#                 if len(file_bytes) == 0:
#                     raise HTTPException(status_code=400, detail="Decoded file is empty")
#             except base64.binascii.Error:
#                 file_result["status"] = "error"
#                 file_result["error"] = "Invalid base64 string"
#                 results.append(file_result)
#                 continue

#             content_type = attachment.type

#             try:
#                 if content_type in ["image/jpeg", "image/png", "image/jpg", "image/gif", "image/tiff", "image/bmp"]:
#                     extracted_texts = [extract_text_and_boxes_from_image(file_bytes)]
#                     extracted_images = [file_bytes]
#                 elif content_type == "application/pdf":
#                     extracted_texts, extracted_images = extract_text_from_pdf(file_bytes)
#                 elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                     extracted_texts = extract_text_from_docx(file_bytes)
#                     extracted_images = []
#                 else:
#                     raise HTTPException(status_code=400, detail="Unsupported file type")
#             except HTTPException as http_err:
#                 file_result["status"] = "error"
#                 file_result["error"] = str(http_err.detail)
#                 results.append(file_result)
#                 continue
#             except Exception as e:
#                 file_result["status"] = "error"
#                 file_result["error"] = f"Failed to extract text: {e}"
#                 results.append(file_result)
#                 continue 

#             if not extracted_texts and not extracted_images:
#                 file_result["status"] = "error"
#                 file_result["error"] = "Failed to extract text from the file"
#                 results.append(file_result)
#                 continue

#             prompt = (
#                 "ReceiptDate (dd-mm-yyyy), Merchant, ExpenseCategory, PaymentMethod (method of payment), "
#                 "ReferenceNumber, TotalAmount (Don't show Currency), CurrencyCode, TaxAmount, "
#                 "in a json format, strictly don't include any other text and it should return like {"
#                 "\"ReceiptDate\": \"\", \"Merchant\": \"\", \"ExpenseCategory\": \"\", \"PaymentMethod\": \"\", "
#                 "\"ReferenceNumber\": \"\", \"TotalAmount\": \"\", \"CurrencyCode\": \"\", \"TaxAmount\": \"\"} without any newline characters."
#             )

#             page_results = []

#             for text in extracted_texts:
#                 try:
#                     structured_data = process_text_with_cohere(text, prompt)
#                     structured_data = clean_json_string(structured_data)

#                     # Log the cleaned structured data for debugging
#                     print("Cleaned structured data:", structured_data)

#                     # Check if the response is valid JSON
#                     if not structured_data.strip():
#                         raise ValueError("Empty JSON response received")

#                     structured_data_json = json.loads(structured_data)

#                     # Filter the JSON response to ensure it only contains the expected keys
#                     expected_keys = ["ReceiptDate", "Merchant", "ExpenseCategory", "PaymentMethod", "ReferenceNumber", "TotalAmount", "CurrencyCode", "TaxAmount"]
#                     filtered_data = {key: structured_data_json.get(key, "") for key in expected_keys}

#                     page_results.append(filtered_data)
#                 except ValueError as ve:
#                     page_results.append({"status": "error", "error": f"Empty JSON response received: {ve}"})
#                 except json.JSONDecodeError as e:
#                     page_results.append({"status": "error", "error": f"Failed to parse JSON response: {e}"})
#                 except Exception as e:
#                     page_results.append({"status": "error", "error": f"Failed to process with Cohere: {e}"})

#             # Process extracted images
#             for image_bytes in extracted_images:
#                 try:
#                     image_text = extract_text_and_boxes_from_image(image_bytes)
#                     structured_data = process_text_with_cohere(image_text, prompt)
#                     structured_data = clean_json_string(structured_data)

#                     # Log the cleaned structured data for debugging
#                     print("Cleaned structured data from image:", structured_data)

#                     # Check if the response is valid JSON
#                     if not structured_data.strip():
#                         raise ValueError("Empty")

#                     structured_data_json = json.loads(structured_data)

#                     # Filter the JSON response to ensure it only contains the expected keys
#                     expected_keys = ["ReceiptDate", "Merchant", "ExpenseCategory", "PaymentMethod", "ReferenceNumber", "TotalAmount", "CurrencyCode", "TaxAmount"]
#                     filtered_data = {key: structured_data_json.get(key, "") for key in expected_keys}

#                     page_results.append(filtered_data)
#                 except ValueError as ve:
#                     page_results.append({"status": "error", "error": f"Empty: {ve}"})
#                 except json.JSONDecodeError as e:
#                     page_results.append({"status": "error", "error": f"Failed to parse JSON response: {e}"})
#                 except Exception as e:
#                     page_results.append({"status": "error", "error": f"Failed to process: {e}"})

#             file_result["extracted_data"] = page_results
#             results.append(file_result)

#         end_time = time.time()  # Record end time
#         processing_time = end_time - start_time

#         return JSONResponse(content={"results": results, "processing_time": processing_time})

#     except HTTPException as http_err:
#         raise http_err
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"An error occurred: {e}")






import streamlit as st
from PIL import Image
from paddleocr import PaddleOCR, draw_ocr
import fitz  # PyMuPDF
from docx import Document
import os
from dotenv import load_dotenv
import numpy as np
from langchain.llms import Ollama
from langchain_community.llms import Ollama
from langchain.prompts import PromptTemplate
import time
import re
# Load environment variables
load_dotenv()

# Initialize PaddleOCR
ocr = PaddleOCR(use_angle_cls=True, lang='en')

# Initialize Ollama with hypothetical llama3 model
ollama = Ollama(base_url="http://localhost:11434", model="llama3")

def extract_text_and_boxes_from_image(image):
    try:
        if image.mode in ["RGBA", "P"]:
            image = image.convert("RGB")
        image_array = np.array(image, dtype=np.uint8)
        result = ocr.ocr(image_array, cls=True)
        
        if not result or not result[0]:
            st.error("No text detected in the image.")
            return None, None, None

        boxes = [line[0] for line in result[0]]
        texts = [line[1][0] for line in result[0]]
        return texts, boxes, result[0]
    except Exception as e:
        st.error(f"An error occurred while extracting text from the image: {e}")
        return None, None, None

def extract_text_from_pdf(uploaded_file):
    try:
        pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text()
        return text
    except Exception as e:
        st.error(f"An error occurred while extracting text from the PDF: {e}")
        return None

def extract_text_from_docx(uploaded_file):
    try:
        doc = Document(uploaded_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"An error occurred while extracting text from the DOCX: {e}")
        return None

def detect_currency(text):
    currency_patterns = {
        'USD': r'\$',
        'EUR': r'€',
        'GBP': r'£',
        'JPY': r'¥',
        'CNY': r'¥',
        'INR': r'₹',
        'AUD': r'A\$',
        'CAD': r'C\$',
        'CHF': r'CHF',
        'SEK': r'kr',
        'NOK': r'kr',
        'DKK': r'kr',
        'NZD': r'NZ\$',
        'SGD': r'S\$',
        'HKD': r'HK\$',
        'ZAR': r'R'
    }
    
    for currency_code, pattern in currency_patterns.items():
        if re.search(pattern, text):
            return currency_code
    
    # If no symbol is found, look for currency codes
    currency_codes = list(currency_patterns.keys())
    for code in currency_codes:
        if re.search(r'\b' + code + r'\b', text):
            return code
    
    return None

def process_text_with_llama3(prompt, text):
    try:
        detected_currency = detect_currency(text)
        
        prompt_template = PromptTemplate(
            input_variables=["prompt", "text", "detected_currency"],
            template="""
            {prompt}

            Text: {text}

            Detected Currency: {detected_currency}

            Please extract the following header-level fields:
            {{
              ReceiptDate: '',
              MerchantName: '',
              ExpenseCategory: '',
              PaymentMethod: '',
              ReferenceNumber: '',
              TotalAmountTrans: '',
              IsReimbursable: '',
              CurrencyCode: '{detected_currency}',
              TaxGroup: '',
              TaxAmount: '',
              Description: ''
            }}
            If any detail is not available, use 'null' or 'no' as appropriate. Do not include any additional text or line-item details.
            Ensure the CurrencyCode is correctly set based on the detected currency. If no currency was detected, use 'null' for CurrencyCode.
            """
        )
        
        formatted_prompt = prompt_template.format(
            prompt=prompt, 
            text=text, 
            detected_currency=detected_currency if detected_currency else 'null'
        )
        
        start_time = time.time()
        response = ollama(formatted_prompt)
        end_time = time.time()
        
        processing_time = end_time - start_time
        return response, processing_time
    except Exception as e:
        st.error(f"An error occurred while processing with LLaMA3: {e}")
        return None, None

# Streamlit app configuration
st.set_page_config(page_title="Header-Level Text Extraction from Images and Documents")
st.header("Header-Level Text Extraction from Images and Documents")

# File uploader for images and documents
uploaded_file = st.file_uploader("Choose an image or document...", type=["jpg", "jpeg", "png", "pdf", "docx"])

if uploaded_file is not None:
    file_type = uploaded_file.type
    if file_type in ["image/jpeg", "image/png", "image/jpg"]:
        image = Image.open(uploaded_file)
        if image:
            st.image(image, caption="Uploaded Image.", use_column_width=True)

            with st.spinner('Extracting text from image...'):
                extracted_text, boxes, results = extract_text_and_boxes_from_image(image)
                if extracted_text:
                    st.subheader("Extracted Text")
                    st.write("\n".join(extracted_text))

                    # Draw bounding boxes on the image
                    font_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'  # Adjust this path to a valid TTF font file on your system
                    if not os.path.exists(font_path):
                        st.error("Font file not found. Please ensure the font file exists at the specified path.")
                    else:
                        image_with_boxes = np.array(image.convert('RGB'))
                        image_with_boxes = draw_ocr(image_with_boxes, boxes, extracted_text, font_path=font_path)
                        image_with_boxes = Image.fromarray(image_with_boxes)
                        st.image(image_with_boxes, caption="Image with Bounding Boxes.", use_column_width=True)

                    # Prompt input after extraction
                    user_prompt = st.text_area("Enter your prompt for LLaMA3 processing:", """
                    Extract the following header-level details in a JSON format:
                    {
                        ReceiptDate: 'dd-mm-yyyy',
                        MerchantName: 'Merchant',
                        ExpenseCategory: 'ExpenseCategory',
                        PaymentMethod: 'PaymentMethod',
                        ReferenceNumber: 'ReferenceNumber',
                        TotalAmountTrans: TotalAmount (numeric value without currency),
                        CurrencyCode: 'CurrencyCode',
                        TaxGroup: 'TaxGroup',
                        TaxAmount: TaxAmount (numeric value),
                        Description: 'Description'
                    }
                    If any detail is not available, use 'null' or 'no' as appropriate. Do not include any additional text or line-item details.
                    """)

                    if st.button("Process with LLaMA3"):
                        with st.spinner('Processing text with LLaMA3...'):
                            structured_data, processing_time = process_text_with_llama3(user_prompt, "\n".join(extracted_text))
                            if structured_data:
                                st.subheader("Structured Data from Document")
                                st.write(structured_data)

                                # Show processing time
                                st.write(f"Processing time: {processing_time:.2f} seconds")

                                # Download button
                                st.download_button(
                                    label="Download Extracted Data",
                                    data=structured_data,
                                    file_name="extracted_data.txt",
                                    mime="text/plain"
                                )
                            else:
                                st.error("Failed to process text with LLaMA3.")
                else:
                    st.error("Failed to extract text from the image.")
    elif file_type == "application/pdf":
        with st.spinner('Extracting text from PDF...'):
            extracted_text = extract_text_from_pdf(uploaded_file)
            if extracted_text:
                st.subheader("Extracted Text")
                st.write(extracted_text)

                # Prompt input after extraction
                user_prompt = st.text_area("Enter your prompt for LLaMA3 processing:", """
                Extract the following header-level details in a JSON format:
                {
                    ReceiptDate: 'dd-mm-yyyy',
                    MerchantName: 'Merchant',
                    ExpenseCategory: 'ExpenseCategory',
                    PaymentMethod: 'PaymentMethod',
                    ReferenceNumber: 'ReferenceNumber',
                    TotalAmountTrans: TotalAmount (numeric value without currency),
                    CurrencyCode: 'CurrencyCode',
                    TaxGroup: 'TaxGroup',
                    TaxAmount: TaxAmount (numeric value),
                    Description: 'Description'
                }
                If any detail is not available, use 'null' or 'no' as appropriate. Do not include any additional text or line-item details.
                """)

                if st.button("Process with LLaMA3"):
                    with st.spinner('Processing text with LLaMA3...'):
                        structured_data, processing_time = process_text_with_llama3(user_prompt, extracted_text)
                        if structured_data:
                            st.subheader("Structured Data from Document")
                            st.write(structured_data)

                            # Show processing time
                            st.write(f"Processing time: {processing_time:.2f} seconds")

                            # Download button
                            st.download_button(
                                label="Download Extracted Data",
                                data=structured_data,
                                file_name="extracted_data.txt",
                                mime="text/plain"
                            )
                        else:
                            st.error("Failed to process text with LLaMA3.")
            else:
                st.error("Failed to extract text from the PDF.")
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        with st.spinner('Extracting text from DOCX...'):
            extracted_text = extract_text_from_docx(uploaded_file)
            if extracted_text:
                st.subheader("Extracted Text")
                st.write(extracted_text)

                # Prompt input after extraction
                user_prompt = st.text_area("Enter your prompt for LLaMA3 processing:", """
                Extract the following header-level details in a JSON format:
                {
                    ReceiptDate: 'dd-mm-yyyy',
                    MerchantName: 'Merchant',
                    ExpenseCategory: 'ExpenseCategory',
                    PaymentMethod: 'PaymentMethod',
                    ReferenceNumber: 'ReferenceNumber',
                    TotalAmountTrans: TotalAmount (numeric value without currency),
                    CurrencyCode: 'CurrencyCode',
                    TaxGroup: 'TaxGroup',
                    TaxAmount: TaxAmount (numeric value),
                    Description: 'Description'
                }
                If any detail is not available, use 'null' or 'no' as appropriate. Do not include any additional text or line-item details.
                """)

                if st.button("Process with LLaMA3"):
                    with st.spinner('Processing text with LLaMA3...'):
                        structured_data, processing_time = process_text_with_llama3(user_prompt, extracted_text)
                        if structured_data:
                            st.subheader("Structured Data from Document")
                            st.write(structured_data)

                            # Show processing time
                            st.write(f"Processing time: {processing_time:.2f} seconds")

                            # Download button
                            st.download_button(
                                label="Download Extracted Data",
                                data=structured_data,
                                file_name="extracted_data.txt",
                                mime="text/plain"
                            )
                        else:
                            st.error("Failed to process text with LLaMA3.")
            else:
                st.error("Failed to extract text from the DOCX.")
else:
    st.info("Please upload an image or document file to get started.")




